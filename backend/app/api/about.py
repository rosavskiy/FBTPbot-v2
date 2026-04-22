"""API страницы «О программе» с журналом доработок."""

from __future__ import annotations

import io
from datetime import date, datetime

from docx import Document
from fastapi import APIRouter, Depends, HTTPException, Query
from fastapi.responses import StreamingResponse
from pydantic import BaseModel, Field
from sqlalchemy import func, select
from sqlalchemy.ext.asyncio import AsyncSession

from app.api.admin_auth import log_action, require_role
from app.config import SARATOV_TZ
from app.database.models import AdminUser, ProgressNote
from app.database.models import get_db as get_admin_db

router = APIRouter(prefix="/api/about", tags=["about"])

_editor = Depends(require_role("superadmin", "admin"))


class ProgressNoteSummary(BaseModel):
    id: int
    progress_date: str
    title: str | None
    content_preview: str
    updated_by: str | None
    updated_at: str | None


class ProgressNoteDetail(BaseModel):
    id: int
    progress_date: str
    title: str | None
    content: str
    created_by: str | None
    updated_by: str | None
    created_at: str | None
    updated_at: str | None


class ProgressNotesListResponse(BaseModel):
    total: int
    items: list[ProgressNoteSummary]


class ProgressNoteUpsertRequest(BaseModel):
    progress_date: date
    title: str | None = Field(None, max_length=255)
    content: str = Field("", max_length=50000)


def _to_summary(note: ProgressNote) -> ProgressNoteSummary:
    preview = (note.content or "").strip().replace("\n", " ")
    preview = preview[:140] + ("..." if len(preview) > 140 else "")
    return ProgressNoteSummary(
        id=note.id,
        progress_date=note.progress_date.isoformat(),
        title=note.title,
        content_preview=preview,
        updated_by=note.updated_by,
        updated_at=note.updated_at.isoformat() if note.updated_at else None,
    )


def _to_detail(note: ProgressNote) -> ProgressNoteDetail:
    return ProgressNoteDetail(
        id=note.id,
        progress_date=note.progress_date.isoformat(),
        title=note.title,
        content=note.content,
        created_by=note.created_by,
        updated_by=note.updated_by,
        created_at=note.created_at.isoformat() if note.created_at else None,
        updated_at=note.updated_at.isoformat() if note.updated_at else None,
    )


def _build_export_filename(note: ProgressNote) -> str:
    return f"about_{note.progress_date.isoformat()}.docx"


def _append_content(document: Document, content: str) -> None:
    lines = [line.rstrip() for line in (content or "").splitlines()]
    if not any(lines):
        document.add_paragraph("Запись пуста.")
        return

    for line in lines:
        stripped = line.strip()
        if not stripped:
            document.add_paragraph("")
            continue
        if stripped.startswith("- "):
            document.add_paragraph(stripped[2:].strip(), style="List Bullet")
            continue
        if stripped.startswith("• "):
            document.add_paragraph(stripped[2:].strip(), style="List Bullet")
            continue
        document.add_paragraph(stripped)


def _build_export_docx(note: ProgressNote) -> io.BytesIO:
    document = Document()
    document.add_heading("О программе", level=1)
    document.add_heading(note.progress_date.strftime("%d.%m.%Y"), level=2)
    if note.title:
        document.add_paragraph(note.title)
    _append_content(document, note.content)
    if note.updated_by:
        document.add_paragraph(f"Последнее обновление: {note.updated_by}")

    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer


def _build_export_range_docx(
    notes: list[ProgressNote],
    date_from: date | None,
    date_to: date | None,
) -> io.BytesIO:
    document = Document()
    document.add_heading("О программе — экспорт за период", level=1)

    parts: list[str] = []
    if date_from:
        parts.append(f"с {date_from.strftime('%d.%m.%Y')}")
    if date_to:
        parts.append(f"по {date_to.strftime('%d.%m.%Y')}")
    if parts:
        document.add_paragraph(" ".join(parts))
    document.add_paragraph(f"Записей в выгрузке: {len(notes)}")

    for note in notes:
        document.add_heading(note.progress_date.strftime("%d.%m.%Y"), level=2)
        if note.title:
            p = document.add_paragraph()
            p.add_run(note.title).bold = True
        _append_content(document, note.content)
        if note.updated_by:
            document.add_paragraph(f"Редактор: {note.updated_by}")

    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer


@router.get("", response_model=ProgressNotesListResponse)
async def list_progress_notes(
    date_from: date | None = Query(None),
    date_to: date | None = Query(None),
    user: AdminUser = _editor,
    db: AsyncSession = Depends(get_admin_db),
):
    query = select(ProgressNote)
    count_query = select(func.count()).select_from(ProgressNote)

    if date_from is not None:
        query = query.where(ProgressNote.progress_date >= date_from)
        count_query = count_query.where(ProgressNote.progress_date >= date_from)
    if date_to is not None:
        query = query.where(ProgressNote.progress_date <= date_to)
        count_query = count_query.where(ProgressNote.progress_date <= date_to)

    query = query.order_by(ProgressNote.progress_date.desc())
    total_result = await db.execute(count_query)
    result = await db.execute(query)
    items = result.scalars().all()

    return ProgressNotesListResponse(
        total=total_result.scalar() or 0,
        items=[_to_summary(item) for item in items],
    )


@router.get("/export-range")
async def export_progress_notes_range(
    date_from: date | None = Query(None),
    date_to: date | None = Query(None),
    user: AdminUser = _editor,
    db: AsyncSession = Depends(get_admin_db),
):
    q = select(ProgressNote).order_by(ProgressNote.progress_date)
    if date_from is not None:
        q = q.where(ProgressNote.progress_date >= date_from)
    if date_to is not None:
        q = q.where(ProgressNote.progress_date <= date_to)

    result = await db.execute(q)
    notes = list(result.scalars().all())

    if not notes:
        raise HTTPException(status_code=404, detail="Записей за указанный период не найдено")

    await log_action(
        db,
        user_id=user.id,
        username=user.username,
        action="export_range",
        entity_type="progress_note",
        entity_id="range",
        entity_name=f"{date_from or '*'}..{date_to or '*'}",
    )

    from_str = date_from.isoformat() if date_from else "all"
    to_str = date_to.isoformat() if date_to else "all"
    filename = f"about_{from_str}_{to_str}.docx"
    headers = {"Content-Disposition": f'attachment; filename="{filename}"'}
    return StreamingResponse(
        _build_export_range_docx(notes, date_from, date_to),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers=headers,
    )


@router.get("/{note_id}", response_model=ProgressNoteDetail)
async def get_progress_note(
    note_id: int,
    user: AdminUser = _editor,
    db: AsyncSession = Depends(get_admin_db),
):
    result = await db.execute(select(ProgressNote).where(ProgressNote.id == note_id))
    note = result.scalar_one_or_none()
    if note is None:
        raise HTTPException(status_code=404, detail="Запись не найдена")
    return _to_detail(note)


@router.get("/by-date/{progress_date}", response_model=ProgressNoteDetail)
async def get_progress_note_by_date(
    progress_date: date,
    user: AdminUser = _editor,
    db: AsyncSession = Depends(get_admin_db),
):
    result = await db.execute(select(ProgressNote).where(ProgressNote.progress_date == progress_date))
    note = result.scalar_one_or_none()
    if note is None:
        raise HTTPException(status_code=404, detail="Запись не найдена")
    return _to_detail(note)


@router.put("", response_model=ProgressNoteDetail)
async def upsert_progress_note(
    payload: ProgressNoteUpsertRequest,
    user: AdminUser = _editor,
    db: AsyncSession = Depends(get_admin_db),
):
    result = await db.execute(select(ProgressNote).where(ProgressNote.progress_date == payload.progress_date))
    note = result.scalar_one_or_none()
    now = datetime.now(SARATOV_TZ)

    if note is None:
        note = ProgressNote(
            progress_date=payload.progress_date,
            title=(payload.title or "").strip() or None,
            content=payload.content,
            created_by=user.username,
            updated_by=user.username,
        )
        db.add(note)
        action = "create"
    else:
        note.title = (payload.title or "").strip() or None
        note.content = payload.content
        note.updated_by = user.username
        note.updated_at = now
        action = "update"

    await db.commit()
    await db.refresh(note)

    await log_action(
        db,
        user_id=user.id,
        username=user.username,
        action=action,
        entity_type="progress_note",
        entity_id=str(note.id),
        entity_name=note.progress_date.isoformat(),
        details=note.title or "Журнал доработок",
    )

    return _to_detail(note)


@router.delete("/{note_id}")
async def delete_progress_note(
    note_id: int,
    user: AdminUser = _editor,
    db: AsyncSession = Depends(get_admin_db),
):
    result = await db.execute(select(ProgressNote).where(ProgressNote.id == note_id))
    note = result.scalar_one_or_none()
    if note is None:
        raise HTTPException(status_code=404, detail="Запись не найдена")

    note_date = note.progress_date.isoformat()
    await db.delete(note)
    await db.commit()

    await log_action(
        db,
        user_id=user.id,
        username=user.username,
        action="delete",
        entity_type="progress_note",
        entity_id=str(note_id),
        entity_name=note_date,
    )

    return {"status": "deleted", "id": note_id}


@router.get("/{note_id}/export")
async def export_progress_note(
    note_id: int,
    user: AdminUser = _editor,
    db: AsyncSession = Depends(get_admin_db),
):
    result = await db.execute(select(ProgressNote).where(ProgressNote.id == note_id))
    note = result.scalar_one_or_none()
    if note is None:
        raise HTTPException(status_code=404, detail="Запись не найдена")

    await log_action(
        db,
        user_id=user.id,
        username=user.username,
        action="export",
        entity_type="progress_note",
        entity_id=str(note.id),
        entity_name=note.progress_date.isoformat(),
    )

    filename = _build_export_filename(note)
    headers = {"Content-Disposition": f'attachment; filename="{filename}"'}
    return StreamingResponse(
        _build_export_docx(note),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers=headers,
    )
